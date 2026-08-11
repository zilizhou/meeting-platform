#!/usr/bin/env python3
"""Build the formal Mingde Tongshu user manual from its Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = "1F4E79"
DEEP_BLUE = "17365D"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
RED = "9C0006"
GOLD = "7F6000"
CONTENT_DXA = 9026  # A4, 2.2 cm left/right margins


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, color=None, font="Noto Sans CJK SC"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (28, DEEP_BLUE, 0, 10),
        "Subtitle": (14, MUTED, 0, 8),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (11.5, DEEP_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans CJK SC")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.95)
        style.paragraph_format.first_line_indent = Cm(-0.45)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_inline(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9.5, color=RED, font="Consolas")
        else:
            clean = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", part)
            run = paragraph.add_run(clean)
            set_run_font(run)


def add_markdown_table(doc, rows):
    if len(rows) < 2:
        return
    data = [rows[0]] + rows[2:]
    cols = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=cols)
    table.style = "Table Grid"
    widths = [CONTENT_DXA // cols] * cols
    widths[-1] += CONTENT_DXA - sum(widths)
    set_table_geometry(table, widths)
    for i, row in enumerate(data):
        for j in range(cols):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if j < len(row):
                add_inline(p, row[j].strip())
            for run in p.runs:
                set_run_font(run, size=9.2, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif j == 0 and cols <= 3:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_picture(doc, source_dir: Path, alt: str, rel_path: str):
    image = source_dir / rel_path
    if not image.exists():
        p = doc.add_paragraph()
        add_inline(p, f"[图片缺失：{alt}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Cm(15.8))
    p.paragraph_format.keep_with_next = True


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("明 德 同 枢")
    set_run_font(r, size=30, bold=True, color=DEEP_BLUE)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("曲阜师范大学二级学院双会议一体化管理系统")
    set_run_font(r, size=17, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("使 用 说 明 书")
    set_run_font(r, size=22, bold=True, color=DEEP_BLUE)
    p.paragraph_format.space_after = Pt(90)
    for text in ("版本：V1.0", "编制日期：2026年8月", "适用范围：业务端与系统管理端"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=11, color=MUTED)
    doc.add_page_break()


def add_front_matter(doc):
    p = doc.add_paragraph("使用提示", style="Heading 1")
    p = doc.add_paragraph()
    add_inline(p, "本说明书面向学院会议工作人员、领导班子成员、校级管理员和校级查阅人员。系统界面与可执行操作会根据账号角色自动收敛；实际权限以系统后端校验和学校制度为准。")
    p = doc.add_paragraph()
    add_inline(p, "重要：AI 功能只用于摘要、草稿和办理引导，不代替人工审题、决策、签署或办结。涉及内部或涉密材料时，应遵守学校信息安全和保密规定。")
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))

    doc.add_paragraph("目录", style="Heading 1")
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "请在 Word 中右键更新目录")
    doc.add_page_break()


def parse_markdown(doc, md_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []
    skip_manual_toc = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.right_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(8)
                if code_lines and code_lines[0].strip().startswith("flowchart"):
                    flow = "议题征集 → 审题 → 创建会议 → 签到表决 → 形成决议 → 纪要签署 → 督办落实 → 归档"
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_run_font(p.add_run(flow), size=10.5, bold=True, color=DEEP_BLUE)
                else:
                    set_run_font(p.add_run("\n".join(code_lines)), size=8.5, color=DEEP_BLUE, font="Consolas")
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not line or line == "---":
            i += 1
            continue
        if line == "## 目录":
            skip_manual_toc = True
            i += 1
            continue
        if skip_manual_toc:
            if line.startswith("## "):
                skip_manual_toc = False
            else:
                i += 1
                continue
        if line.startswith("# "):
            # The cover supplies the document title.
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="Heading 3")
        elif line.startswith("!["):
            m = re.match(r"!\[([^]]*)\]\(([^)]+)\)", line)
            if m:
                add_picture(doc, md_path.parent, m.group(1), m.group(2))
        elif line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_markdown_table(doc, table_rows)
            continue
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            add_inline(p, line.lstrip("> "))
            for run in p.runs:
                run.font.color.rgb = RGBColor.from_string(GOLD)
        elif line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            set_run_font(p.add_run(line.strip("*")), size=9, color=MUTED)
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(header.add_run("明德同枢 · 系统使用说明书"), size=8.5, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(footer.add_run("曲阜师范大学二级学院双会议一体化管理系统  |  "), size=8, color=MUTED)
        add_field(footer, "PAGE", "1")


def build(source: Path, output: Path):
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    add_cover(doc)
    add_front_matter(doc)
    parse_markdown(doc, source)
    core = doc.core_properties
    core.title = "明德同枢——曲阜师范大学二级学院双会议一体化管理系统使用说明书"
    core.subject = "系统操作、角色权限、会议流程、校级监管"
    core.author = "曲阜师范大学"
    core.keywords = "明德同枢,党委会,党政联席会,使用说明书"
    core.comments = "根据现行系统功能编制"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_user_manual.py SOURCE.md OUTPUT.docx")
    build(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
